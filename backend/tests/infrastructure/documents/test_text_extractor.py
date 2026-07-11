import pytest
from pathlib import Path

from app.infrastructure.documents.text_extractor import (
    TextExtractionError,
    UnsupportedFileTypeError,
    extract_text,
)


async def test_extract_txt_reads_utf8_content(tmp_path: Path) -> None:
    file_path = tmp_path / "notes.txt"
    file_path.write_text("Scope 1 emissions: 120 tCO2e", encoding="utf-8")

    result = await extract_text(file_path)

    assert result == "Scope 1 emissions: 120 tCO2e"


async def test_extract_docx_reads_paragraphs(tmp_path: Path) -> None:
    from docx import Document

    file_path = tmp_path / "report.docx"
    document = Document()
    document.add_paragraph("Sustainability Report 2026")
    document.add_paragraph("Total renewable energy usage: 42%")
    document.save(str(file_path))

    result = await extract_text(file_path)

    assert "Sustainability Report 2026" in result
    assert "Total renewable energy usage: 42%" in result


async def test_extract_xlsx_reads_rows_tab_separated(tmp_path: Path) -> None:
    from openpyxl import Workbook

    file_path = tmp_path / "metrics.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    assert sheet is not None
    sheet.append(["Metric", "Value"])
    sheet.append(["CO2e", 100])
    workbook.save(str(file_path))

    result = await extract_text(file_path)

    assert "Metric\tValue" in result
    assert "CO2e\t100" in result


async def test_extract_pdf_reads_page_text(tmp_path: Path) -> None:
    import pymupdf

    file_path = tmp_path / "policy.pdf"
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "ESG Policy Statement")
    doc.save(str(file_path))
    doc.close()

    result = await extract_text(file_path)

    assert "ESG Policy Statement" in result


async def test_extract_text_raises_for_unsupported_extension(tmp_path: Path) -> None:
    file_path = tmp_path / "archive.zip"
    file_path.write_bytes(b"not a real zip")

    with pytest.raises(UnsupportedFileTypeError):
        await extract_text(file_path)


async def test_extract_text_raises_for_corrupt_file(tmp_path: Path) -> None:
    file_path = tmp_path / "broken.docx"
    file_path.write_bytes(b"this is not a valid docx file")

    with pytest.raises(TextExtractionError):
        await extract_text(file_path)
